from pathlib import Path
from tempfile import NamedTemporaryFile
from typing import Any

from services.llm_client import LLMClient, LLMError, with_ai_metadata


class TemplateParseError(ValueError):
    pass


def extract_template_file(upload) -> tuple[str, str]:
    suffix = Path(upload.name).suffix.lower()
    # 限制上传文件大小（默认 20MB），避免超大文件解析拖垮进程
    if getattr(upload, "size", 0) > 20 * 1024 * 1024:
        raise TemplateParseError("文件过大，请上传小于 20MB 的 Word/Excel 文档。")

    import os
    tmp_path = None
    try:
        # 用 delete=False 让第三方库（python-docx/openpyxl）能可靠地按路径重新打开文件
        with NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
            for chunk in upload.chunks():
                tmp.write(chunk)
            tmp.flush()
            tmp_path = tmp.name
        if suffix == ".docx":
            return upload.name, extract_docx_text(tmp_path)
        if suffix == ".xlsx":
            return upload.name, extract_xlsx_text(tmp_path)
        raise TemplateParseError("仅支持 .docx 和 .xlsx 文件；旧版 .doc/.xls 请先另存为新版格式。")
    finally:
        if tmp_path:
            try:
                os.unlink(tmp_path)
            except OSError:
                pass


def extract_docx_text(path: str) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise TemplateParseError("缺少 python-docx 依赖，无法解析 Word 文件。") from exc

    document = Document(path)
    lines = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            lines.append(text)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                lines.append(" | ".join(cells))
    return "\n".join(lines).strip()


def extract_xlsx_text(path: str) -> str:
    try:
        from openpyxl import load_workbook
    except ImportError as exc:
        raise TemplateParseError("缺少 openpyxl 依赖，无法解析 Excel 文件。") from exc

    workbook = load_workbook(path, read_only=True, data_only=True)
    lines = []
    for sheet in workbook.worksheets:
        lines.append(f"工作表：{sheet.title}")
        for row in sheet.iter_rows(values_only=True):
            values = [str(value).strip() for value in row if value not in (None, "")]
            if values:
                lines.append(" | ".join(values))
    return "\n".join(lines).strip()


def extract_position_template(text: str, llm: LLMClient | None = None) -> dict[str, Any]:
    fallback = fallback_extract(text)
    if not llm:
        return with_ai_metadata(fallback, "local_fallback")

    system_prompt = (
        "你是招聘岗位模板智能分析助手。只返回JSON对象，不要返回Markdown。"
        "JSON必须包含 name, department_name, job_level, scenario, description, responsibilities, requirements, technical_tags, keywords。"
        "job_level只能是 junior, middle, senior, expert 之一；无法判断时用 middle。"
        "responsibilities/requirements/technical_tags/keywords 必须是字符串数组。"
        "technical_tags 偏技术栈/工具/框架（如 Python、RAG、LangChain），keywords 偏业务/能力关键词（如 大模型应用、Agent、项目落地）。"
        "若文档存在更多内容，请注意本文档可能被截断，优先抽取前文出现的字段。"
    )
    user_prompt = f"""
请**分析**下面的岗位文档内容，根据岗位职责、任职要求等信息**生成**一个标准、简洁的岗位名称（name）。
注意：不要直接使用文档中工作表的标题或占位名称（如"AI人才画像-xx岗位（xx部门）（请填写）"这类模板标题），
而是基于实际的岗位描述生成准确的岗位名称，例如"Python后端开发工程师"、"大模型算法工程师"、"前端架构师"等。

如果文档中有明确的部门名称，请提取到 department_name。

不要编造文档中没有依据的公司或薪资信息。

文档文本（若内容被截断，以前文为准）：
{text[:12000]}
"""
    try:
        payload = llm.json_completion(system_prompt, user_prompt)
        return with_ai_metadata(merge_payload(payload, fallback), "llm")
    except LLMError as exc:
        fallback["_ai_error"] = str(exc)
        return with_ai_metadata(fallback, "local_fallback_after_error")


def merge_payload(payload: dict[str, Any], fallback: dict[str, Any]) -> dict[str, Any]:
    return {
        "name": clean_string(payload.get("name")) or fallback["name"],
        "department_name": clean_string(payload.get("department_name")) or fallback["department_name"],
        "job_level": normalize_level(payload.get("job_level")) or fallback["job_level"],
        "scenario": clean_string(payload.get("scenario")) or fallback["scenario"],
        "description": clean_string(payload.get("description")) or fallback["description"],
        "responsibilities": clean_list(payload.get("responsibilities")) or fallback["responsibilities"],
        "requirements": clean_list(payload.get("requirements")) or fallback["requirements"],
        "technical_tags": clean_list(payload.get("technical_tags")) or fallback["technical_tags"],
        "keywords": clean_list(payload.get("keywords")) or fallback["keywords"],
    }


def fallback_extract(text: str) -> dict[str, Any]:
    lines = [line.strip().strip("：:") for line in text.splitlines() if line.strip()]
    name = first_after_labels(lines, ["岗位名称", "招聘岗位", "职位名称"]) or guess_name(lines)
    department = first_after_labels(lines, ["所属部门", "用人部门", "部门"]) or ""
    description = section_text(lines, ["岗位描述", "职位描述", "岗位简介"], ["岗位职责", "工作职责", "任职要求", "职位要求", "技能要求"]) or "\n".join(lines[:3])
    responsibilities = section_items(lines, ["岗位职责", "工作职责", "职责描述"], ["任职要求", "职位要求", "技能要求", "岗位要求"])
    requirements = section_items(lines, ["任职要求", "职位要求", "岗位要求"], ["技能标签", "技术方向", "岗位关键词"])
    technical_tags, keywords = guess_tags_split(text)
    return {
        "name": name or "未命名岗位模板",
        "department_name": department,
        "job_level": "middle",
        "scenario": "社会招聘",
        "description": description[:1200],
        "responsibilities": responsibilities[:12],
        "requirements": requirements[:12],
        "technical_tags": technical_tags,
        "keywords": keywords,
    }


def first_after_labels(lines: list[str], labels: list[str]) -> str:
    for line in lines:
        for label in labels:
            if line.startswith(label):
                value = line.replace(label, "", 1).strip(" ：:|-")
                if value:
                    return value
    return ""


def guess_name(lines: list[str]) -> str:
    # 过滤掉模板占位名称
    _placeholder_patterns = ["请填写", "xx", "XXX", "模板", "AI人才画像"]

    def is_placeholder(name: str) -> bool:
        return any(p in name for p in _placeholder_patterns)

    for line in lines[:8]:
        if line.startswith(("岗位描述", "职位描述", "岗位职责", "任职要求", "职位要求")):
            continue
        if "工程师" in line or "专家" in line or "经理" in line or "岗位" in line:
            if not is_placeholder(line):
                return line[:128]
    # 没有匹配到明确的岗位名时，尝试取第一行有意义的短文本
    for line in lines[:3]:
        if 2 < len(line) < 60 and not is_placeholder(line):
            return line[:128]
    return lines[0][:128] if lines else ""


def section_text(lines: list[str], starts: list[str], stops: list[str]) -> str:
    items = section_items(lines, starts, stops, keep_sentences=True)
    return "\n".join(items)


def section_items(lines: list[str], starts: list[str], stops: list[str], keep_sentences: bool = False) -> list[str]:
    collecting = False
    items = []
    for line in lines:
        if any(line.startswith(start) for start in starts):
            collecting = True
            value = line.split("：", 1)[-1].split(":", 1)[-1].strip() if ("：" in line or ":" in line) else ""
            if value:
                items.append(clean_bullet(value))
            continue
        if collecting and any(line.startswith(stop) for stop in stops):
            break
        if collecting:
            cleaned = clean_bullet(line)
            if cleaned:
                if keep_sentences or len(cleaned) > 4:
                    items.append(cleaned)
    return items


def guess_tags(text: str) -> list[str]:
    candidates = ["Python", "Java", "Go", "RAG", "Agent", "Prompt", "LangChain", "向量数据库", "知识库", "大模型", "AI安全", "SQL", "Redis", "Django", "FastAPI"]
    return [item for item in candidates if item.lower() in text.lower()]


def guess_tags_split(text: str) -> tuple[list[str], list[str]]:
    """fallback 模式下区分技术栈与业务关键词，避免两者完全重复。"""
    tech_candidates = ["Python", "Java", "Go", "C++", "SQL", "Redis", "Django", "FastAPI", "LangChain", "TensorFlow", "PyTorch", "向量数据库"]
    keyword_candidates = ["RAG", "Agent", "Prompt", "知识库", "大模型", "大模型应用", "AI工程化", "AI安全", "智能体", "Skills", "项目落地", "架构设计"]
    technical_tags = [item for item in tech_candidates if item.lower() in text.lower()][:8]
    keywords = [item for item in keyword_candidates if item.lower() in text.lower()][:8]
    return technical_tags, keywords


def clean_bullet(value: str) -> str:
    return value.strip().lstrip("0123456789.、)）-• ").strip()


def clean_string(value) -> str:
    return str(value).strip() if value not in (None, "") else ""


def clean_list(value) -> list[str]:
    if not isinstance(value, list):
        return []
    return [clean_string(item) for item in value if clean_string(item)]


def normalize_level(value) -> str:
    value = clean_string(value).lower()
    mapping = {
        "junior": "junior",
        "初级": "junior",
        "middle": "middle",
        "中级": "middle",
        "senior": "senior",
        "高级": "senior",
        "expert": "expert",
        "专家": "expert",
    }
    return mapping.get(value, "")
