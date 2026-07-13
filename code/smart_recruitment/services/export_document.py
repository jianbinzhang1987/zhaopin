"""从已落库的招聘评测数据生成可下载文档。

支持两类格式：
- docx  使用 python-docx
- pdf   使用 reportlab（中文用 STSong-Late CID 字体）

所有导出函数返回 tuple[str, bytes]：(文件名, 内容字节)，供视图设置下载响应。
"""
from __future__ import annotations

import io
import re
from typing import Any

# reportlab 中文 CID 字体
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.platypus import HRFlowable, KeepTogether, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from docx import Document
from django.utils import timezone

from apps.recruitment.models import DevelopmentTask, RecruitmentTask

# 注册中文 CID 字体一次即可
_CN_FONT = "STSong-Light"
_CN_FONT_REGISTERED = False


def _ensure_cn_font() -> str:
    global _CN_FONT_REGISTERED
    if not _CN_FONT_REGISTERED:
        try:
            pdfmetrics.registerFont(UnicodeCIDFont(_CN_FONT))
        except Exception:
            # 降级：若注册失败则返回 Helvetica，中文会乱码但避免抛错
            return "Helvetica"
        _CN_FONT_REGISTERED = True
    return _CN_FONT


# ---------------------------------------------------------------------------
# 通用工具
# ---------------------------------------------------------------------------

_FORBIDDEN_CHARS = re.compile(r'[\\/:*?"<>|\r\n\t]+')


def _safe_filename(text: str, suffix: str) -> str:
    """生成适合下载和本地保存的文件名，去除文件系统非法字符。"""
    cleaned = _FORBIDDEN_CHARS.sub("_", text).strip("_. ") or "export"
    cleaned = cleaned[:120]
    return f"{cleaned}.{suffix}"


def _export_version(task: RecruitmentTask, kind: str) -> int:
    """返回当前导出内容的业务版本号。"""
    if kind == "regular":
        question_set = task.regular_question_sets.order_by("-version").first()
        return question_set.version if question_set else 1
    if kind == "development":
        development_task = task.development_tasks.order_by("-version").first()
        return development_task.version if development_task else 1
    # 评测报告目前没有独立版本表，统一从 v1 开始。
    return 1


def _export_type_label(kind: str, with_answers: bool = True) -> str:
    if kind == "regular":
        return "普通题用人部门版" if with_answers else "普通题候选人版"
    return {"development": "现场开发题", "report": "评测报告"}[kind]


def _export_filename(task: RecruitmentTask, kind: str, suffix: str, with_answers: bool = True) -> str:
    """按“题目类型-候选人-YYYYMMDD-v版本”生成统一下载文件名。"""
    type_label = _export_type_label(kind, with_answers)
    candidate_name = (task.candidate.name or "候选人").strip() or "候选人"
    export_date = timezone.localdate().strftime("%Y%m%d")
    version = _export_version(task, kind)
    return _safe_filename(f"{type_label}-{candidate_name}-{export_date}-v{version}", suffix)


def _title_for(task: RecruitmentTask, kind: str, with_answers: bool) -> str:
    label = {"regular": "普通题", "development": "现场开发题", "report": "评测报告"}[kind]
    edition = "部门版" if (kind != "report" and with_answers) else ("候选人版" if kind == "regular" else "")
    edition_label = f"（{edition}）" if edition else ""
    return f"{task.task_name} {label}{edition_label}"


def _kv_rows(task: RecruitmentTask) -> list[list[str]]:
    return [
        ["任务编号", task.task_no or "-"],
        ["岗位", f"{task.position.name}（{task.position.get_job_level_display()}）"],
        ["候选人", task.candidate.name or "-"],
        ["用人部门", task.department.name if task.department else "-"],
        ["技术负责人", task.technical_owner.get_full_name() or task.technical_owner.username if task.technical_owner else "-"],
        ["生成时间", task.updated_at.strftime("%Y-%m-%d %H:%M") if task.updated_at else "-"],
    ]


def _recommendation_display(rec: str | None) -> str:
    mapping = {"strong_yes": "强烈推荐", "yes": "推荐", "hold": "待定", "no": "不推荐"}
    return mapping.get(rec or "", "待生成")


def _questions_for_export(task: RecruitmentTask) -> list[dict[str, Any]]:
    qs = task.regular_question_sets.order_by("-version").first()
    return qs.questions if qs and isinstance(qs.questions, list) else []


def _dev_task_for_export(task: RecruitmentTask) -> DevelopmentTask | None:
    return task.development_tasks.order_by("-version").first()


# ---------------------------------------------------------------------------
# DOCX 生成
# ---------------------------------------------------------------------------

def build_regular_questions_docx(task: RecruitmentTask, with_answers: bool = True) -> tuple[str, bytes]:
    doc = Document()
    doc.add_heading(_title_for(task, "regular", with_answers), level=1)
    _write_meta_table_docx(doc, task)
    questions = _questions_for_export(task)
    if not questions:
        doc.add_paragraph("暂无题目。")
        return _export_filename(task, "regular", "docx", with_answers), _docx_bytes(doc)
    for idx, q in enumerate(questions, 1):
        doc.add_heading(f"{idx}. [{q.get('skill', '-')}] {q.get('content', '-')}", level=2)
        meta = f"题型 {q.get('type', '-')} · 难度 {q.get('difficulty', '-')}"
        doc.add_paragraph(meta)
        if with_answers:
            ref = q.get("reference_answer", "")
            if ref:
                doc.add_paragraph("参考答案：", style="Intense Quote")
                doc.add_paragraph(ref)
            points = q.get("scoring_points", [])
            if points:
                doc.add_paragraph("评分要点：")
                for p in points:
                    doc.add_paragraph(p, style="List Bullet")
    filename = _export_filename(task, "regular", "docx", with_answers)
    return filename, _docx_bytes(doc)


def build_development_task_docx(task: RecruitmentTask) -> tuple[str, bytes]:
    doc = Document()
    doc.add_heading(_title_for(task, "development", True), level=1)
    _write_meta_table_docx(doc, task)
    dev = _dev_task_for_export(task)
    content = (dev.content if dev else {}) or {}
    if not content:
        doc.add_paragraph("暂无现场开发题内容。")
        return _export_filename(task, "development", "docx"), _docx_bytes(doc)
    _section_paragraph(doc, "业务背景", content.get("background"))
    _section_list(doc, "任务要求", content.get("requirements"))
    _section_paragraph(doc, "开发时长", content.get("duration"))
    _section_dict_list(doc, "约束条件", content.get("constraints"))
    _section_list(doc, "交付内容", content.get("deliverables"))
    _section_list(doc, "验收标准", content.get("acceptance_criteria"))
    filename = _export_filename(task, "development", "docx")
    return filename, _docx_bytes(doc)


def build_report_docx(task: RecruitmentTask) -> tuple[str, bytes]:
    doc = Document()
    doc.add_heading(_title_for(task, "report", True), level=1)
    _write_meta_table_docx(doc, task)
    ev = getattr(task, "evaluation", None)
    if not ev:
        doc.add_paragraph("暂无评测报告数据。")
        return _export_filename(task, "report", "docx"), _docx_bytes(doc)
    # 评分卡片
    table = doc.add_table(rows=4, cols=2)
    table.style = "Light Grid Accent 1"
    rows = [
        ("普通题得分", str(ev.regular_score) if ev.regular_score is not None else "-"),
        ("现场开发题得分", str(ev.development_score) if ev.development_score is not None else "-"),
        ("综合得分", str(ev.final_score) if ev.final_score is not None else "-"),
        ("建议结论", _recommendation_display(ev.recommendation)),
    ]
    for i, (k, v) in enumerate(rows):
        table.cell(i, 0).text = k
        table.cell(i, 1).text = v
    _section_paragraph(doc, "优势", ev.strengths)
    _section_paragraph(doc, "风险", ev.risks)
    skills = ev.skill_evaluations or []
    if skills:
        doc.add_heading("能力评价", level=2)
        sk_table = doc.add_table(rows=1, cols=3)
        sk_table.style = "Light Grid Accent 1"
        hdr = sk_table.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text = "能力", "评价", "依据"
        for s in skills:
            row = sk_table.add_row().cells
            row[0].text = str(s.get("skill", "-"))
            row[1].text = str(s.get("level", "-"))
            row[2].text = str(s.get("evidence", "-"))
    if ev.report_markdown:
        doc.add_heading("报告正文", level=2)
        for line in ev.report_markdown.splitlines():
            if line.strip():
                doc.add_paragraph(line)
    filename = _export_filename(task, "report", "docx")
    return filename, _docx_bytes(doc)


def _write_meta_table_docx(doc: Document, task: RecruitmentTask) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.style = "Light List Accent 1"
    for k, v in _kv_rows(task):
        row = table.add_row().cells
        row[0].text = k
        row[1].text = v
    doc.add_paragraph("")


def _section_paragraph(doc: Document, title: str, value: Any) -> None:
    if value in (None, "", []):
        return
    doc.add_heading(title, level=2)
    doc.add_paragraph(str(value))


def _section_list(doc: Document, title: str, items: Any) -> None:
    if not items:
        return
    doc.add_heading(title, level=2)
    if isinstance(items, str):
        doc.add_paragraph(items)
        return
    for it in items:
        doc.add_paragraph(str(it), style="List Bullet")


def _section_dict_list(doc: Document, title: str, data: Any) -> None:
    if not data:
        return
    doc.add_heading(title, level=2)
    if isinstance(data, dict):
        for k, v in data.items():
            doc.add_paragraph(f"{k}：{v}", style="List Bullet")
    elif isinstance(data, list):
        for it in data:
            doc.add_paragraph(str(it), style="List Bullet")
    else:
        doc.add_paragraph(str(data))


def _docx_bytes(doc: Document) -> bytes:
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# PDF 生成
# ---------------------------------------------------------------------------

def _pdf_styles():
    font = _ensure_cn_font()
    styles = getSampleStyleSheet()
    normal = ParagraphStyle(
        "cn_normal",
        parent=styles["Normal"],
        fontName=font,
        fontSize=10.5,
        leading=18,
        textColor=colors.HexColor("#1f2937"),
        alignment=TA_LEFT,
    )
    title = ParagraphStyle(
        "cn_title",
        parent=styles["Title"],
        fontName=font,
        fontSize=18,
        leading=26,
        alignment=TA_CENTER,
        textColor=colors.HexColor("#111827"),
        spaceAfter=12,
    )
    subtitle = ParagraphStyle(
        "cn_subtitle",
        parent=normal,
        fontSize=9,
        leading=13,
        textColor=colors.HexColor("#667085"),
        alignment=TA_CENTER,
        spaceAfter=10,
    )
    h2 = ParagraphStyle(
        "cn_h2",
        parent=normal,
        fontSize=13,
        leading=19,
        textColor=colors.HexColor("#111827"),
        spaceBefore=12,
        spaceAfter=7,
    )
    h3 = ParagraphStyle(
        "cn_h3",
        parent=normal,
        fontSize=11.5,
        leading=17,
        textColor=colors.HexColor("#111827"),
        spaceBefore=6,
        spaceAfter=5,
    )
    meta = ParagraphStyle("cn_meta", parent=normal, fontSize=9, leading=14, textColor=colors.HexColor("#475467"))
    label = ParagraphStyle("cn_label", parent=meta, textColor=colors.HexColor("#667085"))
    bullet = ParagraphStyle("cn_bullet", parent=normal, leftIndent=12, firstLineIndent=-8, spaceAfter=3)
    answer = ParagraphStyle(
        "cn_answer",
        parent=normal,
        fontSize=10,
        leading=17,
        leftIndent=0,
        rightIndent=0,
        textColor=colors.HexColor("#344054"),
    )
    return {
        "normal": normal,
        "title": title,
        "subtitle": subtitle,
        "h2": h2,
        "h3": h3,
        "meta": meta,
        "label": label,
        "bullet": bullet,
        "answer": answer,
        "font": font,
    }


def _pdf_doc(buf: io.BytesIO, title: str) -> SimpleDocTemplate:
    return SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=18 * mm,
        rightMargin=18 * mm,
        topMargin=18 * mm,
        bottomMargin=18 * mm,
        title=title,
        author="智能招聘评测系统",
    )


def _build_pdf(doc: SimpleDocTemplate, story: list, styles) -> None:
    title = doc.title or "智能招聘评测系统"

    def draw_page(canvas, document):
        canvas.saveState()
        canvas.setFont(styles["font"], 8)
        canvas.setFillColor(colors.HexColor("#98A2B3"))
        canvas.drawString(document.leftMargin, A4[1] - 10 * mm, title[:42])
        canvas.drawRightString(A4[0] - document.rightMargin, 10 * mm, f"第 {document.page} 页")
        canvas.setStrokeColor(colors.HexColor("#EAECF0"))
        canvas.line(document.leftMargin, A4[1] - 13 * mm, A4[0] - document.rightMargin, A4[1] - 13 * mm)
        canvas.restoreState()

    doc.build(story, onFirstPage=draw_page, onLaterPages=draw_page)


def build_regular_questions_pdf(task: RecruitmentTask, with_answers: bool = True) -> tuple[str, bytes]:
    buf = io.BytesIO()
    styles = _pdf_styles()
    doc = _pdf_doc(buf, _title_for(task, "regular", with_answers))
    story = _pdf_meta_block(task, styles, kind="regular", with_answers=with_answers)
    questions = _questions_for_export(task)
    if not questions:
        story.append(Paragraph("暂无题目。", styles["normal"]))
    for idx, q in enumerate(questions, 1):
        _pdf_question(story, styles, idx, q, with_answers)
    _build_pdf(doc, story, styles)
    return _export_filename(task, "regular", "pdf", with_answers), buf.getvalue()


def build_development_task_pdf(task: RecruitmentTask) -> tuple[str, bytes]:
    buf = io.BytesIO()
    styles = _pdf_styles()
    doc = _pdf_doc(buf, _title_for(task, "development", True))
    story = _pdf_meta_block(task, styles, kind="development", with_answers=True)
    dev = _dev_task_for_export(task)
    content = (dev.content if dev else {}) or {}
    if not content:
        story.append(Paragraph("暂无现场开发题内容。", styles["normal"]))
    else:
        _pdf_section(story, styles, "业务背景", content.get("background"))
        _pdf_list(story, styles, "任务要求", content.get("requirements"))
        _pdf_section(story, styles, "开发时长", content.get("duration"))
        if content.get("constraints"):
            data = content["constraints"]
            if isinstance(data, dict):
                text = "；".join(f"{k}：{v}" for k, v in data.items())
            else:
                text = "；".join(str(x) for x in data)
            _pdf_section(story, styles, "约束条件", text)
        _pdf_list(story, styles, "交付内容", content.get("deliverables"))
        _pdf_list(story, styles, "验收标准", content.get("acceptance_criteria"))
    _build_pdf(doc, story, styles)
    return _export_filename(task, "development", "pdf"), buf.getvalue()


def build_report_pdf(task: RecruitmentTask) -> tuple[str, bytes]:
    buf = io.BytesIO()
    styles = _pdf_styles()
    doc = _pdf_doc(buf, _title_for(task, "report", True))
    story = _pdf_meta_block(task, styles, kind="report", with_answers=True)
    ev = getattr(task, "evaluation", None)
    if not ev:
        story.append(Paragraph("暂无评测报告数据。", styles["normal"]))
        _build_pdf(doc, story, styles)
        return _export_filename(task, "report", "pdf"), buf.getvalue()
    score_table = _pdf_kv_table(
        [
            ["普通题得分", _score(ev.regular_score)],
            ["现场开发题得分", _score(ev.development_score)],
            ["综合得分", _score(ev.final_score)],
            ["建议结论", _recommendation_display(ev.recommendation)],
        ],
        styles,
        label_width=44 * mm,
        value_width=52 * mm,
    )
    story.append(score_table)
    story.append(Spacer(1, 10))
    _pdf_section(story, styles, "优势", ev.strengths)
    _pdf_section(story, styles, "风险", ev.risks)
    skills = ev.skill_evaluations or []
    if skills:
        story.append(Paragraph("能力评价", styles["h2"]))
        rows = [["能力", "评价", "依据"]] + [
            [_p(s.get("skill"), styles["normal"]), _p(s.get("level"), styles["normal"]), _p(s.get("evidence"), styles["normal"])] for s in skills
        ]
        sk = Table(rows, colWidths=[40 * mm, 30 * mm, 90 * mm])
        sk.setStyle(_pdf_table_style(styles, header=True))
        story.append(sk)
    if ev.report_markdown:
        _pdf_section(story, styles, "报告正文", ev.report_markdown)
    _build_pdf(doc, story, styles)
    return _export_filename(task, "report", "pdf"), buf.getvalue()


def _pdf_meta_block(task: RecruitmentTask, styles, kind: str = "regular", with_answers: bool = True) -> list:
    title = _title_for(task, kind, with_answers)
    story = [
        Paragraph(title, styles["title"]),
        Paragraph("智能招聘评测系统导出文档", styles["subtitle"]),
        HRFlowable(width="100%", thickness=0.8, color=colors.HexColor("#EAECF0")),
        Spacer(1, 8),
    ]
    meta = _pdf_kv_table(_kv_rows(task), styles, label_width=32 * mm, value_width=132 * mm)
    story.append(meta)
    story.append(Spacer(1, 12))
    return story


def _pdf_question(story, styles, idx: int, q: dict[str, Any], with_answers: bool) -> None:
    title = f"{idx}. [{_s(q.get('skill') or '-')}] {_s(q.get('content') or '-')}"
    rows = [
        [Paragraph(title, styles["h3"])],
        [Paragraph(f"题型：{_s(q.get('type') or '-')}　难度：{_s(q.get('difficulty') or '-')}", styles["meta"])],
    ]
    if with_answers:
        ref = q.get("reference_answer")
        if ref:
            rows.append([[Paragraph("<b>参考答案</b>", styles["label"]), Spacer(1, 3), Paragraph(_s(ref), styles["answer"])]])
        points = q.get("scoring_points") or []
        if points:
            rows.append([[Paragraph("<b>评分要点</b>", styles["label"]), Spacer(1, 3), _pdf_bullets(points, styles)]])
    card = Table(rows, colWidths=[164 * mm], splitByRow=1)
    card.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F8FAFC")),
        ("BOX", (0, 0), (-1, -1), 0.6, colors.HexColor("#D0D5DD")),
        ("LINEBELOW", (0, 0), (-1, 0), 0.6, colors.HexColor("#EAECF0")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 10),
        ("RIGHTPADDING", (0, 0), (-1, -1), 10),
        ("TOPPADDING", (0, 0), (-1, -1), 7),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
    ]))
    story.append(KeepTogether([card, Spacer(1, 8)]))


def _pdf_section(story, styles, title: str, value: Any) -> None:
    if value in (None, "", []):
        return
    story.append(Paragraph(title, styles["h2"]))
    story.append(Paragraph(_s(value), styles["normal"]))
    story.append(Spacer(1, 6))


def _pdf_list(story, styles, title: str, items: Any) -> None:
    if not items:
        return
    story.append(Paragraph(title, styles["h2"]))
    if isinstance(items, str):
        story.append(Paragraph(_s(items), styles["normal"]))
    else:
        story.append(_pdf_bullets(items, styles))
    story.append(Spacer(1, 6))


def _pdf_bullets(items: Any, styles) -> Table:
    rows = []
    for item in items:
        rows.append([Paragraph("-", styles["normal"]), Paragraph(_s(item), styles["normal"])])
    bullets = Table(rows, colWidths=[6 * mm, 150 * mm])
    bullets.setStyle(TableStyle([
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 0),
        ("RIGHTPADDING", (0, 0), (-1, -1), 3),
        ("TOPPADDING", (0, 0), (-1, -1), 1),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 1),
        ("TEXTCOLOR", (0, 0), (0, -1), colors.HexColor("#0F6BFF")),
    ]))
    return bullets


def _pdf_kv_table(rows: list[list[Any]], styles, label_width: float, value_width: float) -> Table:
    data = [[_p(k, styles["label"]), _p(v, styles["normal"])] for k, v in rows]
    table = Table(data, colWidths=[label_width, value_width], hAlign="LEFT")
    table.setStyle(_pdf_table_style(styles))
    return table


def _pdf_table_style(styles, header: bool = False) -> TableStyle:
    commands = [
        ("FONTNAME", (0, 0), (-1, -1), styles["font"]),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#D0D5DD")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
        ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#F8FAFC")),
    ]
    if header:
        commands.extend([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#EEF4FF")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#1D4ED8")),
        ])
    return TableStyle(commands)


def _p(value: Any, style: ParagraphStyle) -> Paragraph:
    return Paragraph(_s(value), style)


def _score(value) -> str:
    return "-" if value in (None, "") else str(value)


def _s(value: Any) -> str:
    """段落文本安全化：转字符串并替换易破坏 reportlab 的尖括号。"""
    if value is None:
        return ""
    text = _clean_export_text(str(value))
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _clean_export_text(text: str) -> str:
    """清理模型生成内容里常见的 Markdown 标记，避免 PDF 中露出格式噪声。"""
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"^\s{0,3}#{1,6}\s*", "", text, flags=re.MULTILINE)
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"`([^`]*)`", r"\1", text)
    text = re.sub(r"^\s*[-*]\s+", "- ", text, flags=re.MULTILINE)
    return text.strip()
